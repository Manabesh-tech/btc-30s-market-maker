from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_PATH = ROOT / "btc_orderbook_imbalance_skew_report_2026-05-27.docx"

TOP5_BEST = {
    "kind": "notional or raw size",
    "depth": 5,
    "threshold": 0.5,
    "persist_seconds": 3.0,
    "payouts": "70/90",
    "today_30_before": -2518.115167,
    "today_30_after": 2889.317833,
    "today_30_improvement": 5407.432999999981,
    "today_30_orders": 12341,
    "today_30_triggered": 3482,
    "today_1m_before": -1184.9125980000001,
    "today_1m_after": -712.7595980000003,
    "today_1m_improvement": 472.1529999999998,
    "today_1m_orders": 1391,
    "today_1m_triggered": 302,
    "last3_30_before": 32221.157513000002,
    "last3_30_after": 52568.181513000105,
    "last3_30_improvement": 20347.0240000001,
    "last3_30_orders": 51964,
    "last3_30_triggered": 13910,
    "last3_1m_before": 5275.085985,
    "last3_1m_after": 6664.007985000001,
    "last3_1m_improvement": 1388.9220000000016,
    "last3_1m_orders": 7458,
    "last3_1m_triggered": 1657,
    "last7_30_before": 49322.709613,
    "last7_30_after": 80939.93861300038,
    "last7_30_improvement": 31617.229000000378,
    "last7_30_orders": 81515,
    "last7_30_triggered": 22486,
    "last7_1m_before": 6618.205385,
    "last7_1m_after": 8481.116385000004,
    "last7_1m_improvement": 1862.911000000004,
    "last7_1m_orders": 12679,
    "last7_1m_triggered": 2845,
}

CONSERVATIVE = {
    "kind": "top 10 notional",
    "depth": 10,
    "threshold": 0.7,
    "persist_seconds": 3.0,
    "payouts": "70/90",
    "today_30_after": -212.875167,
    "today_30_improvement": 2305.24,
    "today_1m_after": -972.885598,
    "today_1m_improvement": 212.02700000000016,
    "last3_30_after": 40629.915513,
    "last3_30_improvement": 8408.757999999998,
    "last3_1m_after": 5839.831985,
    "last3_1m_improvement": 564.7460000000001,
    "last7_30_after": 62787.497613,
    "last7_30_improvement": 13464.788,
    "last7_1m_after": 7350.054385,
    "last7_1m_improvement": 731.8490000000002,
}

TOP_RANKED = [
    ["Top 5", "Notional", "0.5", "31,617.23", "1,862.91", "5,407.43", "472.15"],
    ["Top 5", "Raw size", "0.5", "31,617.23", "1,862.91", "5,407.43", "472.15"],
    ["Top 10", "Notional", "0.5", "31,391.77", "1,837.23", "5,268.28", "463.40"],
    ["Top 10", "Raw size", "0.5", "31,391.77", "1,837.23", "5,268.28", "463.40"],
    ["Top 25", "Notional", "0.5", "29,795.27", "1,687.85", "4,949.57", "438.79"],
    ["Top 25", "Raw size", "0.5", "29,783.38", "1,687.85", "4,949.57", "438.79"],
]


def fmt_money(x: float) -> str:
    return f"{x:,.2f}"


def fmt_pct(x: float) -> str:
    return f"{x:.1f}%"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_text(cell, text: str, bold: bool = False, color: str | None = None, size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], fill: str = "EAF1F8") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, h in enumerate(headers):
        set_text(table.rows[0].cells[i], h, bold=True)
        shade(table.rows[0].cells[i], fill)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_text(cells[i], value)
    doc.add_paragraph("")


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10)


def trigger_share(triggered: int, total: int) -> str:
    return fmt_pct(triggered / total * 100)


def build_doc() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(11.5)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    rt = title.add_run("BTC Orderbook Imbalance Skew Report")
    rt.bold = True
    rt.font.name = "Arial"
    rt.font.size = Pt(22)
    rt.font.color.rgb = RGBColor(24, 73, 122)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(8)
    rs = sub.add_run(
        "Backtest through May 27, 2026 16:08 SGT. Focus: imbalance-based skewing for BTC 30s and BTC 1m."
    )
    rs.font.name = "Arial"
    rs.font.size = Pt(10.5)
    rs.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_heading("Executive Summary", level=1)
    add_bullet(
        doc,
        "I tested Binance Futures BTCUSDT orderbook imbalance using Tardis historical depth data and applied a 70/90 payout skew when the book pressure was one-sided and persistent.",
    )
    add_bullet(
        doc,
        "Among the tested constructions, the maximum savings came from a top-5 orderbook imbalance rule with threshold 0.5 and 3-second persistence.",
    )
    add_bullet(
        doc,
        "For BTCUSDT, top-5 notional imbalance and top-5 raw-size imbalance were effectively identical in savings. This is expected because the first five price levels are very close in price, so value weighting and size weighting behave almost the same.",
    )
    add_bullet(
        doc,
        f"Over the last 7 days, that maximum-savings rule would have improved BTC 30s platform PnL by {fmt_money(TOP5_BEST['last7_30_improvement'])} and BTC 1m platform PnL by {fmt_money(TOP5_BEST['last7_1m_improvement'])}.",
    )
    add_bullet(
        doc,
        "The tradeoff is aggressiveness: the max-savings rule triggers on a much larger share of orders than the more conservative top-10 / 0.7 rule."
    )

    doc.add_heading("What Orderbook Imbalance Means", level=1)
    p = doc.add_paragraph()
    p.add_run("Raw size imbalance: ").bold = True
    p.add_run("compare only quantity resting on the bid side versus the ask side.")
    add_code(doc, "bid_size_topN = Σ(size_bid_i),   ask_size_topN = Σ(size_ask_i)")
    add_code(doc, "raw_size_imbalance = (bid_size_topN - ask_size_topN) / (bid_size_topN + ask_size_topN)")

    p = doc.add_paragraph()
    p.add_run("Notional imbalance: ").bold = True
    p.add_run("weight each level by price x size, so the signal compares resting value, not just raw contracts.")
    add_code(doc, "bid_notional_topN = Σ(price_bid_i × size_bid_i),   ask_notional_topN = Σ(price_ask_i × size_ask_i)")
    add_code(doc, "notional_imbalance = (bid_notional_topN - ask_notional_topN) / (bid_notional_topN + ask_notional_topN)")

    doc.add_paragraph(
        "Interpretation for both measures is the same: +1 means the book is heavily bid, -1 means heavily offered, and 0 means balanced."
    )
    doc.add_paragraph(
        "Because BTCUSDT top-of-book prices are nearly identical across the first few levels, notional and raw-size imbalance are almost the same signal for top-5 depth. That is exactly what the backtest results show."
    )

    doc.add_heading("Trading Rule Applied In The Backtest", level=1)
    add_table(
        doc,
        ["Component", "Rule"],
        [
            ["Base payouts", "80 / 80"],
            ["Skewed payouts", "70 / 90"],
            ["Momentum side", "The side aligned with the orderbook imbalance sign"],
            ["Trigger sign", "Book imbalance > threshold for UP, < -threshold for DOWN"],
            ["Persistence rule", "Same sign must persist for at least 3 seconds"],
            ["Universe", "BTC only, durations 30s and 1m"],
            ["PnL basis", "Raw order-level platform PnL before and after skew, same basis on both sides"],
        ],
    )

    doc.add_heading("Data And Methodology", level=1)
    add_bullet(
        doc,
        "Exchange signal source: Tardis historical Binance Futures BTCUSDT orderbook data.",
    )
    add_bullet(
        doc,
        "Completed UTC days were annotated from Tardis daily book_snapshot_25 files. The still-open UTC day was annotated from Tardis replay so today could be included before the daily export completes.",
    )
    add_bullet(
        doc,
        "Depth constructions tested: top 5, top 10, and top 25 levels.",
    )
    add_bullet(
        doc,
        "Imbalance constructions tested: notional imbalance and raw-size imbalance.",
    )
    add_bullet(
        doc,
        "Thresholds tested: 0.5, 0.6, 0.7, 0.8 with fixed 3-second persistence.",
    )
    add_bullet(
        doc,
        "Backtest assumption: user flow is unchanged after skew. So all savings here are first-order pricing savings, not a behavioral re-optimization.",
    )
    add_bullet(
        doc,
        "Current-day cutoff: May 27, 2026 16:08 SGT. Tardis current-day data had a short live lag, so the analysis was cut to the latest available depth timestamp.",
    )

    doc.add_heading("Which Construction Saves The Most", level=1)
    doc.add_paragraph(
        "The table below ranks the highest-savings constructions from the tested grid. The objective here is maximum savings, not minimum operational disruption."
    )
    add_table(
        doc,
        ["Depth", "Imbalance", "Threshold", "Last 7d 30s Δ", "Last 7d 1m Δ", "Today 30s Δ", "Today 1m Δ"],
        TOP_RANKED,
        fill="D9EAD3",
    )

    doc.add_paragraph(
        "Result: top 5 depth with threshold 0.5 is the best-performing construction. Notional and raw-size are effectively tied at that depth."
    )

    doc.add_heading("Maximum-Savings Rule", level=1)
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Depth", "Top 5 levels"],
            ["Imbalance type", "Notional or raw size (effectively identical here)"],
            ["Threshold", "0.5"],
            ["Persistence", "3 seconds"],
            ["Payout skew", "70 / 90"],
            ["Interpretation", "Skew the side aligned with strong orderbook pressure down to 70, and the other side up to 90"],
        ],
        fill="FCE4D6",
    )

    doc.add_heading("Savings Summary For The Maximum-Savings Rule", level=1)
    add_table(
        doc,
        ["Window", "Product", "Without Skew", "With Skew", "Improvement", "Triggered Orders", "Triggered Share"],
        [
            [
                "Today",
                "BTC 30s",
                fmt_money(TOP5_BEST["today_30_before"]),
                fmt_money(TOP5_BEST["today_30_after"]),
                fmt_money(TOP5_BEST["today_30_improvement"]),
                f"{TOP5_BEST['today_30_triggered']:,} / {TOP5_BEST['today_30_orders']:,}",
                trigger_share(TOP5_BEST["today_30_triggered"], TOP5_BEST["today_30_orders"]),
            ],
            [
                "Today",
                "BTC 1m",
                fmt_money(TOP5_BEST["today_1m_before"]),
                fmt_money(TOP5_BEST["today_1m_after"]),
                fmt_money(TOP5_BEST["today_1m_improvement"]),
                f"{TOP5_BEST['today_1m_triggered']:,} / {TOP5_BEST['today_1m_orders']:,}",
                trigger_share(TOP5_BEST["today_1m_triggered"], TOP5_BEST["today_1m_orders"]),
            ],
            [
                "Last 3 days",
                "BTC 30s",
                fmt_money(TOP5_BEST["last3_30_before"]),
                fmt_money(TOP5_BEST["last3_30_after"]),
                fmt_money(TOP5_BEST["last3_30_improvement"]),
                f"{TOP5_BEST['last3_30_triggered']:,} / {TOP5_BEST['last3_30_orders']:,}",
                trigger_share(TOP5_BEST["last3_30_triggered"], TOP5_BEST["last3_30_orders"]),
            ],
            [
                "Last 3 days",
                "BTC 1m",
                fmt_money(TOP5_BEST["last3_1m_before"]),
                fmt_money(TOP5_BEST["last3_1m_after"]),
                fmt_money(TOP5_BEST["last3_1m_improvement"]),
                f"{TOP5_BEST['last3_1m_triggered']:,} / {TOP5_BEST['last3_1m_orders']:,}",
                trigger_share(TOP5_BEST["last3_1m_triggered"], TOP5_BEST["last3_1m_orders"]),
            ],
            [
                "Last 7 days",
                "BTC 30s",
                fmt_money(TOP5_BEST["last7_30_before"]),
                fmt_money(TOP5_BEST["last7_30_after"]),
                fmt_money(TOP5_BEST["last7_30_improvement"]),
                f"{TOP5_BEST['last7_30_triggered']:,} / {TOP5_BEST['last7_30_orders']:,}",
                trigger_share(TOP5_BEST["last7_30_triggered"], TOP5_BEST["last7_30_orders"]),
            ],
            [
                "Last 7 days",
                "BTC 1m",
                fmt_money(TOP5_BEST["last7_1m_before"]),
                fmt_money(TOP5_BEST["last7_1m_after"]),
                fmt_money(TOP5_BEST["last7_1m_improvement"]),
                f"{TOP5_BEST['last7_1m_triggered']:,} / {TOP5_BEST['last7_1m_orders']:,}",
                trigger_share(TOP5_BEST["last7_1m_triggered"], TOP5_BEST["last7_1m_orders"]),
            ],
        ],
        fill="FFF2CC",
    )

    doc.add_heading("Comparison Against The More Conservative Rule", level=1)
    doc.add_paragraph(
        "The top-5 / 0.5 rule maximizes savings, but it is much more aggressive than the previously discussed top-10 / 0.7 rule. The comparison below shows the difference."
    )
    add_table(
        doc,
        ["Window", "Product", "Max-Savings Rule Δ", "Conservative Rule Δ", "Difference"],
        [
            ["Today", "BTC 30s", fmt_money(TOP5_BEST["today_30_improvement"]), fmt_money(CONSERVATIVE["today_30_improvement"]), fmt_money(TOP5_BEST["today_30_improvement"] - CONSERVATIVE["today_30_improvement"])],
            ["Today", "BTC 1m", fmt_money(TOP5_BEST["today_1m_improvement"]), fmt_money(CONSERVATIVE["today_1m_improvement"]), fmt_money(TOP5_BEST["today_1m_improvement"] - CONSERVATIVE["today_1m_improvement"])],
            ["Last 3 days", "BTC 30s", fmt_money(TOP5_BEST["last3_30_improvement"]), fmt_money(CONSERVATIVE["last3_30_improvement"]), fmt_money(TOP5_BEST["last3_30_improvement"] - CONSERVATIVE["last3_30_improvement"])],
            ["Last 3 days", "BTC 1m", fmt_money(TOP5_BEST["last3_1m_improvement"]), fmt_money(CONSERVATIVE["last3_1m_improvement"]), fmt_money(TOP5_BEST["last3_1m_improvement"] - CONSERVATIVE["last3_1m_improvement"])],
            ["Last 7 days", "BTC 30s", fmt_money(TOP5_BEST["last7_30_improvement"]), fmt_money(CONSERVATIVE["last7_30_improvement"]), fmt_money(TOP5_BEST["last7_30_improvement"] - CONSERVATIVE["last7_30_improvement"])],
            ["Last 7 days", "BTC 1m", fmt_money(TOP5_BEST["last7_1m_improvement"]), fmt_money(CONSERVATIVE["last7_1m_improvement"]), fmt_money(TOP5_BEST["last7_1m_improvement"] - CONSERVATIVE["last7_1m_improvement"])],
        ],
        fill="D9EAD3",
    )

    doc.add_paragraph(
        "The practical implication is straightforward: top-5 / 0.5 is the best money-saving rule in the tested set, but it skews a much larger slice of flow. Top-10 / 0.7 sacrifices a large amount of savings in exchange for much less intervention."
    )

    doc.add_heading("Interpretation", level=1)
    add_bullet(
        doc,
        "Shallower book pressure is more useful than deeper book pressure for these ultra-short products. Top 5 outperformed top 10 and top 25.",
    )
    add_bullet(
        doc,
        "Lower threshold 0.5 saved the most money because it triggers earlier and captures more of the one-sided pressure users appear to trade with.",
    )
    add_bullet(
        doc,
        "The 30s product is much more sensitive to this signal than the 1m product. The skew still helps 1m, but the incremental savings are far smaller.",
    )
    add_bullet(
        doc,
        "Notional and raw-size versions are almost identical at top 5 for BTCUSDT. In practice, that means the implementation can choose the simpler internal representation without sacrificing much performance.",
    )

    doc.add_heading("Caveats", level=1)
    add_bullet(
        doc,
        "All results use raw order-level platform PnL rather than rebate-adjusted official platform PnL, because the available DB user does not have access to the cashbook rebate table. The before/after comparison remains consistent because the same basis is used on both sides.",
    )
    add_bullet(
        doc,
        "This is a first-order pricing counterfactual. It assumes users continue to trade the same size and side even after seeing the skewed quote. Realized live savings may be lower if user behavior changes.",
    )
    add_bullet(
        doc,
        "The current-day cut stops at the latest Tardis timestamp available during the run, so today is not an end-of-day figure.",
    )

    doc.add_heading("Files Used", level=1)
    add_bullet(doc, "imbalance_construction_backtest_2026-05-27.json")
    add_bullet(doc, "btc_skew_summary_30s_1m_2026-05-27.json")
    add_bullet(doc, "analyze_imbalance_constructions.py")
    add_bullet(doc, "annotate_orders_tardis_replay_multi.mjs")
    add_bullet(doc, "tardis_datasets_book25/*")

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
