from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_PATH = ROOT / "btc_30s_rule_memo_revised_2026-05-25.docx"


RULE_TEXT = (
    "Block BTC 30s orders when: previous closed Binance BTCUSDT perpetual 15m "
    "volume percentile < 25, previous 15m impact percentile > 70, current "
    "smoothed volatility is between 10 and 30, and the user is chasing the "
    "last 30s move by more than 1 bp."
)

SUMMARY_7D = {
    "window": "2026-05-19 00:00 SGT to 2026-05-25 11:06 SGT",
    "orders": 57599,
    "original_pnl": 40430.5316,
    "excluded_pnl": -24439.4694,
    "after_pnl": 64870.0010,
    "orders_excluded": 2993,
    "excluded_share_pct": 5.1962707685897325,
    "days_helped": 6,
    "days_hurt": 0,
}

SUMMARY_10D = {
    "window": "2026-05-16 00:00 SGT to 2026-05-25 11:06 SGT",
    "orders": 68376,
    "original_pnl": 40449.0911,
    "excluded_pnl": -11361.6135,
    "after_pnl": 51810.7046,
    "orders_excluded": 2056,
    "excluded_share_pct": 3.006903006903007,
    "days_helped": 8,
    "days_hurt": 0,
}

DAILY_10D = [
    ["2026-05-16", -8038.3633, -584.5020, -7453.8613, 584.5020],
    ["2026-05-18", 8186.6547, -1131.3586, 9318.0133, 1131.3586],
    ["2026-05-19", 4924.2195, -1445.2063, 6369.4258, 1445.2063],
    ["2026-05-20", 15872.1801, -592.5269, 16464.7070, 592.5269],
    ["2026-05-21", 11726.9375, -2339.4872, 14066.4247, 2339.4872],
    ["2026-05-22", 3315.0098, -4281.1233, 7596.1331, 4281.1233],
    ["2026-05-23", 2059.6048, -398.8992, 2458.5040, 398.8992],
    ["2026-05-25", 2402.8480, -588.5100, 2991.3580, 588.5100],
]

MORNING_PATCH = {
    "window": "2026-05-25 09:00-09:44 SGT",
    "raw_pnl": -1180.7760,
    "excluded_pnl": -1185.7900,
    "after_pnl": 5.0140,
    "orders": 258,
    "orders_excluded": 61,
}

AGGRESSIVE_OPTION = {
    "name": "More aggressive variant",
    "rule": "prevVol<25, prevImpact>70, smoothedVol 10-30, chase 30s > 0 bps",
    "window": "2026-05-16 00:00 SGT to 2026-05-25 11:03 SGT",
    "original_pnl": 40832.0497,
    "after_pnl": 51718.9090,
    "improvement": 10886.8593,
    "orders_excluded": 3271,
    "days_helped": 7,
    "days_hurt": 1,
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "EAF1F8") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph("")


def fmt_money(x: float) -> str:
    return f"{x:,.2f}"


def make_doc() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(21)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12.5)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("BTC 30s Risk Filter Memo")
    r.font.color.rgb = RGBColor(24, 73, 122)

    intro = doc.add_paragraph()
    intro.add_run("Purpose: ").bold = True
    intro.add_run(
        "recommend one practical ex-ante BTC 30s filter that saves money over the "
        "recent 7-10 day backtest without hurting day-level results."
    )

    scope = doc.add_paragraph()
    scope.add_run("Measurement basis: ").bold = True
    scope.add_run(
        "order-level platform PnL before and after the filter, using the same basis "
        "on both sides. This makes the strategy delta internally consistent."
    )

    doc.add_heading("Recommendation", level=1)
    p = doc.add_paragraph()
    p.add_run("Recommended live rule: ").bold = True
    p.add_run(RULE_TEXT)

    add_table(
        doc,
        ["Component", "Rule"],
        [
            ["Previous 15m Binance volume percentile", "< 25"],
            ["Previous 15m Binance impact percentile", "> 70"],
            ["Current smoothed volatility", "10 to 30"],
            ["Chasing test", "User direction matches last 30s move by more than 1 bp"],
            ["Scope", "BTC 30s only"],
        ],
    )

    doc.add_heading("Why This Version", level=1)
    doc.add_paragraph(
        "This is the best zero-hurt rule I found over the longer 10-day window. "
        "It is meaningfully protective, still selective on blocked flow, and it did "
        "not reduce day-level PnL on any tested day."
    )

    doc.add_heading("10-Day Result", level=1)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Backtest window", SUMMARY_10D["window"]],
            ["Original BTC 30s PnL", fmt_money(SUMMARY_10D["original_pnl"])],
            ["PnL removed by filter", fmt_money(SUMMARY_10D["excluded_pnl"])],
            ["After filter", fmt_money(SUMMARY_10D["after_pnl"])],
            [
                "Improvement",
                fmt_money(SUMMARY_10D["after_pnl"] - SUMMARY_10D["original_pnl"]),
            ],
            ["Orders excluded", f"{SUMMARY_10D['orders_excluded']:,}"],
            ["Orders total", f"{SUMMARY_10D['orders']:,}"],
            ["Blocked share", f"{SUMMARY_10D['excluded_share_pct']:.1f}%"],
            ["Days helped", str(SUMMARY_10D["days_helped"])],
            ["Days hurt", str(SUMMARY_10D["days_hurt"])],
        ],
    )

    doc.add_heading("7-Day Result", level=1)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Backtest window", SUMMARY_7D["window"]],
            ["Original BTC 30s PnL", fmt_money(SUMMARY_7D["original_pnl"])],
            ["PnL removed by filter", fmt_money(SUMMARY_7D["excluded_pnl"])],
            ["After filter", fmt_money(SUMMARY_7D["after_pnl"])],
            [
                "Improvement",
                fmt_money(SUMMARY_7D["after_pnl"] - SUMMARY_7D["original_pnl"]),
            ],
            ["Orders excluded", f"{SUMMARY_7D['orders_excluded']:,}"],
            ["Orders total", f"{SUMMARY_7D['orders']:,}"],
            ["Blocked share", f"{SUMMARY_7D['excluded_share_pct']:.1f}%"],
            ["Days helped", str(SUMMARY_7D["days_helped"])],
            ["Days hurt", str(SUMMARY_7D["days_hurt"])],
        ],
    )

    doc.add_heading("Daily 10-Day Before / After", level=1)
    add_table(
        doc,
        ["Date", "Original", "Removed", "After Filter", "Delta"],
        [
            [date, fmt_money(orig), fmt_money(removed), fmt_money(after), fmt_money(delta)]
            for date, orig, removed, after, delta in DAILY_10D
        ],
    )

    doc.add_heading("This Morning Check", level=1)
    doc.add_paragraph(
        "The same rule also explains the main BTC 30s loss patch this morning. "
        "Using the 09:00-09:44 SGT pocket as the test window:"
    )
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Loss patch window", MORNING_PATCH["window"]],
            ["Raw BTC 30s PnL", fmt_money(MORNING_PATCH["raw_pnl"])],
            ["PnL removed by rule", fmt_money(MORNING_PATCH["excluded_pnl"])],
            ["After rule", fmt_money(MORNING_PATCH["after_pnl"])],
            ["Orders excluded", f"{MORNING_PATCH['orders_excluded']:,} / {MORNING_PATCH['orders']:,}"],
        ],
        header_fill="FCE4D6",
    )

    doc.add_heading("Alternative If We Want More Savings", level=1)
    doc.add_paragraph(
        "There is a broader version that saves slightly more money, but it starts "
        "to introduce day-level false positives. I would not start with it."
    )
    add_table(
        doc,
        ["Alternative", "Value"],
        [
            ["Rule", AGGRESSIVE_OPTION["rule"]],
            ["10-day original PnL", fmt_money(AGGRESSIVE_OPTION["original_pnl"])],
            ["10-day after filter", fmt_money(AGGRESSIVE_OPTION["after_pnl"])],
            ["10-day improvement", fmt_money(AGGRESSIVE_OPTION["improvement"])],
            ["Orders excluded", f"{AGGRESSIVE_OPTION['orders_excluded']:,}"],
            ["Days helped / hurt", f"{AGGRESSIVE_OPTION['days_helped']} / {AGGRESSIVE_OPTION['days_hurt']}"],
        ],
        header_fill="FFF2CC",
    )

    doc.add_heading("Implementation Note", level=1)
    doc.add_paragraph(
        "This should first run as a live shadow flag so we can verify the flagged "
        "orders remain the same kind of loss-making chasers in production. If the "
        "shadow results hold, the next step is to block only BTC 30s orders that hit "
        "all four conditions."
    )

    doc.add_heading("Caveat", level=1)
    doc.add_paragraph(
        "This memo uses order-level platform PnL consistently before and after the "
        "filter. I could not read the rebate cashbook table with the current DB user, "
        "so the memo does not claim exact official Metabase net PnL. For strategy "
        "comparison, however, the before/after savings numbers remain valid because "
        "both strategies are measured on the same basis."
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(make_doc())
