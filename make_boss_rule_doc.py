from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
JSON_PATH = ROOT / "btc_30s_filter_search_7d.json"
OUT_PATH = ROOT / "btc_30s_ex_ante_rule_memo.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        set_cell_shading(table.rows[0].cells[i], "EAF1F8")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    doc.add_paragraph("")


def main() -> None:
    data = json.loads(JSON_PATH.read_text())
    top_name = data["ranked_candidates"][0]["name"]
    detail = next(x for x in data["top_candidate_details"] if x["name"] == top_name)
    original = data["original_pnl"]
    orders = data["orders"]
    excluded_orders = detail["orders_excluded"]
    after = detail["after_pnl"]
    improvement = detail["improvement"]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)

    header = section.header.paragraphs[0]
    header.text = "BTC 30s Ex-Ante Risk Rule Memo"
    header.style = styles["Normal"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(102, 102, 102)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("BTC 30s Ex-Ante Filter Recommendation")
    run.font.color.rgb = RGBColor(24, 73, 122)

    sub = doc.add_paragraph()
    sub.add_run("Prepared for: internal risk / product review\n").bold = True
    sub.add_run("Scope: BTC 30s only, last 7 Singapore days, raw no-rebate PnL backtest")
    sub.paragraph_format.space_after = Pt(12)

    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph()
    p.add_run("Recommended rule: ").bold = True
    p.add_run(
        "block BTC 30s trades when the previous closed Binance BTCUSDT perp 15m bar is low-volume and high-impact, "
        "the current internal smoothed volatility is medium-high, and the user is chasing the last 30-second move."
    )
    p = doc.add_paragraph()
    p.add_run("Exact ex-ante rule: ").bold = True
    p.add_run(
        "exclude the trade if "
        "(1) previous 15m Binance volume percentile < 25, "
        "(2) previous 15m Binance impact percentile > 70, "
        "(3) current smoothed_vol_pct is between 15 and 30, and "
        "(4) the trade direction matches the last 30s BTC move."
    )
    p = doc.add_paragraph()
    p.add_run("7-day result: ").bold = True
    p.add_run(
        f"raw BTC 30s platform PnL improved from {original:,.2f} to {after:,.2f}, "
        f"an uplift of {improvement:,.2f}, while excluding {excluded_orders:,} of {orders:,} orders "
        f"({excluded_orders / orders * 100:.1f}%)."
    )

    doc.add_heading("Recommended Rule", level=1)
    add_table(
        doc,
        ["Component", "Definition", "Threshold"],
        [
            ["Previous 15m volume percentile", "Binance BTCUSDT perp quote-volume percentile using only prior history", "< 25"],
            ["Previous 15m impact percentile", "Impact percentile using only prior history", "> 70"],
            ["Current volatility regime", "Internal BTC 30s smoothed_vol_pct at order entry", "15 to 30"],
            ["Chasing condition", "Trade direction matches last 30s BTC move", "True"],
        ],
    )

    doc.add_heading("Operational Definitions", level=1)
    doc.add_paragraph(
        "1. Previous 15m volume percentile: percentile rank of the previous closed Binance BTCUSDT perpetual 15m quote-volume "
        "compared only against prior historical observations. The implementation first compares against prior bars from the same "
        "15-minute clock slot; if there is not enough slot history, it falls back to all prior 15m bars."
    )
    doc.add_paragraph(
        "2. Previous 15m impact percentile: percentile rank of the previous closed Binance BTCUSDT perpetual 15m impact, where "
        "impact = range_bps / (quote_volume_usd / 1,000,000), again using only prior historical observations."
    )
    doc.add_paragraph(
        "3. Current smoothed volatility: the existing internal 30-second smoothed_vol_pct measure computed from oracle price logs."
    )
    doc.add_paragraph(
        "4. Chasing condition: if the last 30-second BTC return is positive, block UP trades; if the last 30-second BTC return is negative, "
        "block DOWN trades. In shorthand: direction x last_30s_return_bps > 0."
    )

    doc.add_heading("7-Day Evidence", level=1)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Backtest window", f"{data['window_sg_start']} to {data['window_sg_end']}"],
            ["Original raw BTC 30s PnL", f"{original:,.2f}"],
            ["After filter", f"{after:,.2f}"],
            ["Improvement", f"{improvement:,.2f}"],
            ["Orders excluded", f"{excluded_orders:,}"],
            ["Orders total", f"{orders:,}"],
            ["Excluded share", f"{excluded_orders / orders * 100:.1f}%"],
            ["Excluded PnL", f"{detail['excluded_pnl']:,.2f}"],
            ["Positive PnL excluded", f"{detail['positive_excluded_pnl']:,.2f}"],
            ["Negative PnL excluded", f"{detail['negative_excluded_pnl']:,.2f}"],
            ["Days helped", str(detail["days_helped"])],
            ["Days hurt", str(detail["days_hurt"])],
        ],
    )

    doc.add_heading("Daily Before / After", level=1)
    daily_rows = []
    for d in detail["daily"]:
        daily_rows.append(
            [
                d["sg_date"],
                f"{d['original_pnl']:,.2f}",
                f"{d['excluded_pnl']:,.2f}",
                f"{d['after_pnl']:,.2f}",
                f"{d['delta']:,.2f}",
            ]
        )
    add_table(doc, ["Date", "Original PnL", "Excluded PnL", "After Filter", "Delta"], daily_rows)

    doc.add_heading("Why This Rule Is Better", level=1)
    doc.add_paragraph(
        "The earlier window-level liquidity rule was not good enough because it blocked both good and bad periods. "
        "The improvement came from adding the user-behavior condition: only block trades when users are chasing the most recent "
        "30-second move inside a thin 15-minute market regime."
    )
    doc.add_paragraph(
        "This makes the filter much more selective. Instead of blocking all trades in a low-volume period, it targets the specific subset "
        "of trades that are most associated with platform losses."
    )

    doc.add_heading("Implementation Notes", level=1)
    doc.add_paragraph(
        "Apply the rule at order time on BTC 30s only. Evaluate the previous closed Binance BTCUSDT perpetual 15m bar, read the current "
        "internal smoothed_vol_pct, compute the last 30-second BTC return, and reject the order only if all four rule conditions are true."
    )
    doc.add_paragraph(
        "This should be deployed first as a shadow / dry-run monitor before hard blocking, so we can verify that the live flagged trades "
        "continue to match the backtest pattern."
    )

    doc.add_heading("Caveats", level=1)
    doc.add_paragraph(
        "1. This backtest is raw no-rebate PnL because the available DB user does not have SELECT access to public.dex_user_cashbooks. "
        "The direction of the result is still useful, but the exact totals should be rechecked with rebate-adjusted PnL before production sign-off."
    )
    doc.add_paragraph(
        "2. This is an in-sample 7-day test. Before hard enforcement, run the rule as a live shadow monitor and verify out-of-sample behavior."
    )
    doc.add_paragraph(
        "3. Scope here is BTC 30s only. ETH and longer durations should be tested separately before reusing the same thresholds."
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
